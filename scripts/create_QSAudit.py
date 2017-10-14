from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
import os
import get_qlik_sense
import shutil
import sys
import argparse
import datetime


from clint.textui import progress

year = datetime.date.today().year
month = datetime.date.today().month
day = datetime.date.today().day

parser = argparse.ArgumentParser()
parser.add_argument('--server', help='Qlik Sense Server to connect to.')
parser.add_argument('--certs', help='Path to certificates.')
parser.add_argument('--user', help='Username in format domain\\user')
parser.add_argument('--wmi', help='True/False, set to True to collect WMI information from servers. This switch requires windows authentication (username and password)')
parser.add_argument('--password', help='Password of user.')
args = parser.parse_args()

try:
    os.remove("Qlik Sense Site.docx")
except FileNotFoundError:
    pass

shutil.copy('./word_template/Qlik Sense Site.docx', 'QSDoc_{0}_{1}_{2}_{3}.docx'.format(args.server, year, month, day))

document = Document('QSDoc_{0}_{1}_{2}_{3}.docx'.format(args.server, year, month, day))
sections = document.sections
section = sections[0]

def connect():
    """
    Attempt to connect to the Qlik Sense server specified in the arguments.
    """
    # print ('Testing connection..')
    get_qlik_sense.get_about()


def appOwner():
    section = document.add_section()
    document.add_heading('Application Owners', level=1)
    apps= get_qlik_sense.getAppOwners()
    num_of_apps = len(apps)
    table = document.add_table(rows=num_of_apps+1, cols=5)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'App name'
    row.cells[1].text = 'Publish time'
    row.cells[2].text = 'Stream'
    row.cells[3].text = 'Owner userId'
    row.cells[4].text = 'Owner userName'
    for app in range(num_of_apps):
        row = table.rows[app+1]
        row.cells[0].text = str(apps[app][0])
        row.cells[1].text = str(apps[app][1])
        row.cells[2].text = str(apps[app][2])
        row.cells[3].text = str(apps[app][3])
        row.cells[4].text = str(apps[app][4])
    document.add_page_break()
    
def appObjectOwner():
    section = document.add_section()
    document.add_heading('Application Object Owners', level=1)
    apps = get_qlik_sense.getAppObjects()
    num_of_apps = len(apps)
    table = document.add_table(rows=num_of_apps+1, cols=5)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'App name'
    row.cells[1].text = 'Object Type'
    row.cells[2].text = 'Object Name'
    row.cells[3].text = 'Owner userId'
    row.cells[4].text = 'Owner userName'
    for app in range(num_of_apps):
        row = table.rows[app+1]
        row.cells[0].text = str(apps[app][0])
        row.cells[1].text = str(apps[app][1])
        row.cells[2].text = str(apps[app][2])
        row.cells[3].text = str(apps[app][3])
        row.cells[3].text = str(apps[app][4])
    document.add_page_break()

def summary():
    # section = document.add_section()
    document.add_heading('Total number of Authors - Applications, Sheets and Stories', level=1)
    summary = get_qlik_sense.totalUsers()
    table = document.add_table(rows=2, cols=1)
    table.style = 'Grid Table 1 Light Accent 1'
    row = table.rows[0]
    row.cells[0].text = 'Number of Authors'
    row = table.rows[+1]
    row.cells[0].text = str(summary)


def test_connection():
    """
    Test the connection to the central node
    """
    try:
        connect()
    except:
        print ('Server {0} could not be reached, please check connection settings..'.format(args.server))
    else:
        main()

def savedoc():
    """
    Save output to word
    """
    document.save('QSDoc_{0}_{1}_{2}_{3}.docx'.format(args.server, year, month, day))

def main():
    """
    Main
    """
    print('Generating report..')
      
    dispatcher = [summary, appOwner,appObjectOwner,savedoc]
 

    for item in progress.bar(range(len(dispatcher))):

        dispatcher[item]()

test_connection()


